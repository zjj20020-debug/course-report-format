#!/usr/bin/env python3
"""
课程报告生成器 — 哈尔滨工程大学学报格式
用法: python generate_report.py [report_content.json]

依赖: pip install lxml python-docx; pandoc >= 3.x (PATH)
"""
import json, os, sys, subprocess, tempfile, shutil

def check_deps():
    if subprocess.run(['pandoc','--version'], capture_output=True).returncode != 0:
        sys.exit('pandoc not found. Install: winget install pandoc')
    try:
        from lxml import etree
        from docx import Document
    except ImportError as e:
        sys.exit(f'Missing Python package: {e}. Run: pip install lxml python-docx')

check_deps()
from lxml import etree
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Paths ──
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
JSON_PATH  = sys.argv[1] if len(sys.argv) > 1 else os.path.join(os.getcwd(), 'report_content.json')
OUT_DIR    = os.path.dirname(os.path.abspath(JSON_PATH))
OUT_PATH   = os.path.join(OUT_DIR, 'RIS_6G_Report.docx')

W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
M = 'http://schemas.openxmlformats.org/officeDocument/2006/math'
XML_SPACE = '{http://www.w3.org/XML/1998/namespace}space'

# ── Load ──
with open(JSON_PATH, encoding='utf-8') as f:
    data = json.load(f)

# ═══════════════════════════════════════════════
# STEP 1: Build base docx via pipeline approach
# Strategy: run pandoc on a pure-Markdown body (no refs/media),
# then post-process with python-docx + lxml
# ═══════════════════════════════════════════════

tmp = tempfile.mkdtemp(prefix='crf_')
base_out = os.path.join(tmp, 'base.docx')

# 1a. Write flat Markdown (text-only, skip refs and media objects)
md_lines = []
for sec in data['sections']:
    if 'refs' in sec: continue
    md_lines += [f'# {sec["heading"]}', '']
    if 'subsections' in sec:
        for sub in sec['subsections']:
            md_lines += [f'## {sub["subheading"]}', '']
            for p in sub['paragraphs']:
                if isinstance(p, str):
                    md_lines += [p, '']
    if 'paragraphs' in sec:
        for p in sec['paragraphs']:
            if isinstance(p, str):
                md_lines += [p, '']

md_path = os.path.join(tmp, 'body.md')
with open(md_path, 'w', encoding='utf-8') as f:
    f.write('\n'.join(md_lines))

# 1b. pandoc → docx with MathML
pandoc_out = os.path.join(tmp, 'pandoc.docx')
r = subprocess.run([
    'pandoc', md_path, '-o', pandoc_out,
    '--mathml', '-f', 'markdown+tex_math_dollars',
], capture_output=True, text=True)
if r.returncode != 0:
    sys.exit(f'pandoc failed: {r.stderr}')

# 1c. Create structure docx with python-docx (Chinese header + column setup)
doc = Document()

# Default style
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.line_spacing = 1.0

# Section 1: single-column (title → CN keywords)
sec1 = doc.sections[0]
sec1.top_margin    = int(2.6 * 567)  # 2.6cm in EMU ≈ 1474 DXA
sec1.bottom_margin = int(1.6 * 567)
sec1.left_margin   = int(2.0 * 567)
sec1.right_margin  = int(2.0 * 567)

import re
TWIPS_PER_PT = 20

def _indent_twips(chars, font_pt=10.5):
    return int(chars * font_pt * TWIPS_PER_PT)

INDENT2 = _indent_twips(2)

def _sp_element(before=0, after=0, line=None, rule=None):
    sp = OxmlElement('w:spacing')
    if before: sp.set(qn('w:before'), str(int(before * TWIPS_PER_PT)))
    if after:  sp.set(qn('w:after'), str(int(after * TWIPS_PER_PT)))
    if line is not None:
        sp.set(qn('w:line'), str(line))
        if rule: sp.set(qn('w:lineRule'), rule)
    return sp

def _jc_element(val='both'):
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), val)
    return jc

def _ind_element(first=0, left=0, right=0):
    ind = OxmlElement('w:ind')
    if first: ind.set(qn('w:firstLine'), str(first))
    if left:  ind.set(qn('w:left'), str(left))
    if right: ind.set(qn('w:right'), str(right))
    return ind

def _run(font, sz_pt, bold=False):
    r = OxmlElement('w:r')
    rpr = OxmlElement('w:rPr')
    rf = OxmlElement('w:rFonts')
    for a in ['ascii','hAnsi','eastAsia','cs']:
        rf.set(qn(f'w:{a}'), font)
    rpr.append(rf)
    if bold:
        b = OxmlElement('w:b'); rpr.append(b)
        bCs = OxmlElement('w:bCs'); rpr.append(bCs)
    sz = OxmlElement('w:sz'); sz.set(qn('w:val'), str(int(sz_pt * 2))); rpr.append(sz)
    szCs = OxmlElement('w:szCs'); szCs.set(qn('w:val'), str(int(sz_pt * 2))); rpr.append(szCs)
    r.append(rpr)
    return r

def _add_p(after_el, text, font, sz_pt, bold=False, align='both',
           sp_before=0, sp_after=0, line=None, line_rule=None,
           ind_first=0, ind_left=0, ind_right=0):
    """Insert <w:p> after after_el, return new element."""
    p = OxmlElement('w:p')
    ppr = OxmlElement('w:pPr')
    ppr.append(_jc_element(align))
    ppr.append(_sp_element(before=sp_before, after=sp_after, line=line, rule=line_rule))
    if ind_first or ind_left or ind_right:
        ppr.append(_ind_element(first=ind_first, left=ind_left, right=ind_right))
    p.insert(0, ppr)

    r = _run(font, sz_pt, bold)
    t = OxmlElement('w:t'); t.set(qn('xml:space'), 'preserve'); t.text = text
    r.append(t); p.append(r)
    after_el.addnext(p)
    return p

def _add_table(after_el, caption_cn, caption_en, headers, rows):
    """三线表: 4200 DXA wide, top/bottom 1pt, header rule 0.75pt, 小5 宋体."""
    el = _add_p(after_el, caption_cn, '黑体', 9, bold=True, align='center', sp_before=6)
    el = _add_p(el, caption_en, 'Times New Roman', 9, bold=True, align='center')

    ncols = len(headers); col_w = int(4200 / ncols)
    tbl = OxmlElement('w:tbl')
    tblPr = OxmlElement('w:tblPr')
    tw = OxmlElement('w:tblW'); tw.set(qn('w:w'), '4200'); tw.set(qn('w:type'), 'dxa'); tblPr.append(tw)
    tblPr.append(_jc_element('center'))
    borders = OxmlElement('w:tblBorders')
    for edge, sz in [('top','8'),('bottom','8'),('left','0'),('right','0'),('insideH','0'),('insideV','0')]:
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), sz)
        b.set(qn('w:space'), '0'); b.set(qn('w:color'), '000000'); borders.append(b)
    tblPr.append(borders); tbl.append(tblPr)
    tblGrid = OxmlElement('w:tblGrid')
    for _ in range(ncols):
        gc = OxmlElement('w:gridCol'); gc.set(qn('w:w'), str(col_w)); tblGrid.append(gc)
    tbl.append(tblGrid)

    def _tc(col_w, content, bottom_border=False):
        tc = OxmlElement('w:tc'); tcPr = OxmlElement('w:tcPr')
        tcW = OxmlElement('w:tcW'); tcW.set(qn('w:w'), str(col_w)); tcW.set(qn('w:type'), 'dxa'); tcPr.append(tcW)
        if bottom_border:
            tcB = OxmlElement('w:tcBorders')
            bb = OxmlElement('w:bottom'); bb.set(qn('w:val'), 'single'); bb.set(qn('w:sz'), '6')
            bb.set(qn('w:space'), '0'); bb.set(qn('w:color'), '000000'); tcB.append(bb); tcPr.append(tcB)
        tc.insert(0, tcPr)
        p2 = OxmlElement('w:p'); ppr2 = OxmlElement('w:pPr')
        ppr2.append(_jc_element('center')); p2.insert(0, ppr2)
        r2 = _run('宋体', 9, bold=bottom_border)
        t2 = OxmlElement('w:t'); t2.set(qn('xml:space'), 'preserve'); t2.text = content; r2.append(t2)
        p2.append(r2); tc.append(p2)
        return tc

    tr = OxmlElement('w:tr')
    for h in headers: tr.append(_tc(col_w, h, bottom_border=True))
    tbl.append(tr)
    for row in rows:
        tr = OxmlElement('w:tr')
        for cell_text in row: tr.append(_tc(col_w, cell_text))
        tbl.append(tr)
    el.addnext(tbl)
    gap = OxmlElement('w:p'); gap.insert(0, OxmlElement('w:pPr')); tbl.addnext(gap)
    return gap

def _add_figure(after_el, caption_cn, caption_en):
    el = _add_p(after_el, '[ 图片 ]', '宋体', 9, align='center')
    el = _add_p(el, caption_cn, '黑体', 9, bold=True, align='center')
    el = _add_p(el, caption_en, 'Times New Roman', 9, bold=True, align='center')
    return el

# ── Chinese header (Section 1 - single column) ──
body = doc.element.find(f'{{{W}}}body')

# Remove default sectPr from python-docx (we add our own later)
for ch in list(body):
    if ch.tag == f'{{{W}}}sectPr':
        body.remove(ch)

def _new_p():
    """Create a new <w:p> and append to body."""
    p = OxmlElement('w:p')
    ppr = OxmlElement('w:pPr')
    p.insert(0, ppr)
    body.append(p)
    return p

# Title: 二号 黑体 居中
title_el = _new_p()
r_title = _run('黑体', 22, bold=True)
t_title = OxmlElement('w:t'); t_title.set(qn('xml:space'), 'preserve'); t_title.text = data['title']
r_title.append(t_title)
title_el.append(r_title)
title_ppr = title_el.find(f'{{{W}}}pPr')
title_ppr.insert(0, _jc_element('center'))
title_ppr.insert(1, _sp_element(before=60, after=15, line=240))

# Author: 小四 楷体 居中
el = _add_p(title_el, data.get('author', '作者姓名'), '楷体', 12, align='center',
            line=300, line_rule='exact', sp_after=6)

# Affiliation: 小五 楷体 居中
el = _add_p(el, f"({data.get('affiliation', '哈尔滨工程大学 信息与通信工程学院，黑龙江 哈尔滨 150001')})",
            '楷体', 9, align='center', line=300, line_rule='exact', sp_after=6)

# Abstract: 小五 楷体 左右缩进2字符
el = _add_p(el, f"摘  要：{data['abstract']}", '楷体', 9, align='both',
            line=300, line_rule='exact', ind_left=INDENT2, ind_right=INDENT2)

# Keywords: 小五 宋体
el = _add_p(el, f"关键词：{data['keywords']}", '宋体', 9, align='both',
            line=300, line_rule='exact', sp_after=12, ind_left=INDENT2, ind_right=INDENT2)

# ── English header (still in Section 1) ──
if data.get('title_en'):
    el = _add_p(el, data['title_en'], 'Times New Roman', 14, bold=True, align='center',
                line=240, sp_after=15)
    el = _add_p(el, data.get('author_en', 'AUTHOR Name'), 'Times New Roman', 12, align='center',
                line=300, line_rule='exact')
    el = _add_p(el, f"({data.get('affiliation_en', '')})", 'Times New Roman', 9, align='center',
                line=240, sp_after=6)
    el = _add_p(el, f"Abstract: {data.get('abstract_en', '')}", 'Times New Roman', 9, align='both',
                line=300, line_rule='exact', sp_after=15, ind_left=INDENT2, ind_right=INDENT2)
    el = _add_p(el, f"Keywords: {data.get('keywords_en', '')}", 'Times New Roman', 9, align='both',
                line=300, line_rule='exact', sp_after=12, ind_left=INDENT2, ind_right=INDENT2)

# Mark the end of Section 1 with a sectPr inside pPr
sec1_end_el = el

# ═══════════════════════════════════════════════
# STEP 2: Insert pandoc body + Section 2 (two-column)
# ═══════════════════════════════════════════════

b_zip = None
try:
    import zipfile
    b_zip = zipfile.ZipFile(pandoc_out, 'r')
    b_xml = etree.fromstring(b_zip.read('word/document.xml'))
    b_body = b_xml.find(f'{{{W}}}body')

    # Copy pandoc body children (skip sectPr)
    for child in b_body:
        if child.tag == f'{{{W}}}sectPr': continue
        body.append(child)
finally:
    if b_zip: b_zip.close()

# Add continuous section break for 2-column layout
# The last element in body now is the last pandoc paragraph
# We need sectPr right after it (but Python-docx puts sectPr as child of last <w:p>)
# Solution: add sectPr as direct child of <w:body>
sec2_pr = OxmlElement('w:sectPr')
sec2_type = OxmlElement('w:type'); sec2_type.set(qn('w:val'), 'continuous')
sec2_pr.append(sec2_type)
sec2_pg = OxmlElement('w:pgSz'); sec2_pg.set(qn('w:w'), '11906'); sec2_pg.set(qn('w:h'), '16838')
sec2_pr.append(sec2_pg)
sec2_margin = OxmlElement('w:pgMar')
sec2_margin.set(qn('w:top'), '1474'); sec2_margin.set(qn('w:bottom'), '907')
sec2_margin.set(qn('w:left'), '1134'); sec2_margin.set(qn('w:right'), '1134')
sec2_pr.append(sec2_margin)
sec2_cols = OxmlElement('w:cols')
sec2_cols.set(qn('w:num'), '2'); sec2_cols.set(qn('w:space'), '720'); sec2_cols.set(qn('w:equalWidth'), 'true')
sec2_pr.append(sec2_cols)
body.append(sec2_pr)

# ═══════════════════════════════════════════════
# STEP 3: Insert tables and figures
# ═══════════════════════════════════════════════

# Collect media items grouped by subsection heading
media_map = {}  # sub_heading -> [(type, item), ...]
for sec in data['sections']:
    if 'subsections' not in sec: continue
    for sub in sec['subsections']:
        items = [x for x in sub['paragraphs'] if isinstance(x, dict)]
        if items:
            media_map[sub['subheading']] = items

# Build paragraph index map from document body
para_elements = [c for c in body if c.tag == f'{{{W}}}p']
para_texts = [''.join(p.itertext()).strip() for p in para_elements]

def find_para_idx(text_fragment):
    for i, t in enumerate(para_texts):
        if text_fragment in t:
            return i
    return -1

for sub_heading, items in media_map.items():
    start = find_para_idx(sub_heading)
    if start < 0: continue
    # Find end of this subsection
    end = len(para_elements)
    all_headings = []
    for sec in data['sections']:
        if 'subsections' in sec:
            for sub in sec['subsections']:
                all_headings.append(sub['subheading'])
    for other_h in all_headings:
        if other_h == sub_heading: continue
        idx = find_para_idx(other_h)
        if idx > start and idx < end:
            end = idx

    # Find last text paragraph in this range
    last_idx = end - 1
    while last_idx > start:
        txt = para_texts[last_idx]
        if txt and not txt.startswith(sub_heading) and not any(txt.startswith(h) for h in all_headings):
            break
        last_idx -= 1
    if last_idx <= start: last_idx = start

    insert_el = para_elements[last_idx]
    for item in items:
        if item.get('type') == 'table':
            insert_el = _add_table(insert_el, item['caption_cn'], item['caption_en'],
                                   item['headers'], item['rows'])
        elif item.get('type') == 'figure':
            insert_el = _add_figure(insert_el, item['caption_cn'], item['caption_en'])

# ═══════════════════════════════════════════════
# STEP 4: References
# ═══════════════════════════════════════════════

ref_section = None
for sec in data['sections']:
    if 'refs' in sec:
        ref_section = sec; break

if ref_section:
    # Title: 四号黑体, 单倍行距
    el = _add_p(body[-1], ref_section['heading'], '黑体', 14, bold=True, align='left',
                line=240, sp_before=10, sp_after=10)

    for ref in ref_section['refs']:
        if ref['type'] == 'en':
            el = _add_p(el, ref['text'], 'Times New Roman', 9, align='both',
                        line=240, line_rule='exact')
        elif ref['type'] == 'cn':
            el = _add_p(el, ref['cn'], '宋体', 9, align='both',
                        line=240, line_rule='exact')
            el = _add_p(el, ref['en'], 'Times New Roman', 9, align='both',
                        line=240, line_rule='exact')

# ═══════════════════════════════════════════════
# STEP 5: Fix Section 1 sectPr (must be inside last paragraph of section 1)
# ═══════════════════════════════════════════════

# Add sectPr to sec1_end_el's pPr
sec1_ppr = sec1_end_el.find(f'{{{W}}}pPr')
if sec1_ppr is None:
    sec1_ppr = OxmlElement('w:pPr')
    sec1_end_el.insert(0, sec1_ppr)

sec1_sectpr = OxmlElement('w:sectPr')
sec1_pgSz = OxmlElement('w:pgSz'); sec1_pgSz.set(qn('w:w'), '11906'); sec1_pgSz.set(qn('w:h'), '16838')
sec1_sectpr.append(sec1_pgSz)
sec1_pgMar = OxmlElement('w:pgMar')
sec1_pgMar.set(qn('w:top'), '1474'); sec1_pgMar.set(qn('w:bottom'), '907')
sec1_pgMar.set(qn('w:left'), '1134'); sec1_pgMar.set(qn('w:right'), '1134')
sec1_sectpr.append(sec1_pgMar)
sec1_ppr.append(sec1_sectpr)

# ═══════════════════════════════════════════════
# STEP 6: Save & validate
# ═══════════════════════════════════════════════

doc.save(OUT_PATH)

# Quick XML validation
import zipfile as zf
with zf.ZipFile(OUT_PATH, 'r') as z:
    for name in z.namelist():
        if name.endswith('.xml') or name.endswith('.rels'):
            try:
                etree.fromstring(z.read(name).decode('utf-8').encode('utf-8'))
            except Exception as e:
                print(f'WARNING: {name} invalid: {e}')

shutil.rmtree(tmp)
print(f'[OK] {OUT_PATH}')
